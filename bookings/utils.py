from docx import Document


class ReportGenerator:
    def create_report(self, bookings):
        document = Document()
        document.add_heading('Meeting Room Booking Report', 0)

        for booking in bookings:
            document.add_paragraph(f'{booking}')

        document.add_page_break()
        return document
